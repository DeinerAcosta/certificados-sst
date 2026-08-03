"""
FOCA · Generador de Certificado de PRUEBA (v2 — QR flotante estilo CAEM)
=========================================================================

Cambios v2:
  · El QR ahora es una imagen FLOTANTE anclada a la esquina
    inferior derecha de la página (como en el certificado CAEM).
  · No se agregan párrafos extras que puedan empujar a página 2.
  · Texto "Escanee para verificar" también flotante al lado del QR.
"""

from pathlib import Path
import qrcode
from qrcode.constants import ERROR_CORRECT_H
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

# ==== CONFIGURACIÓN DE LA PRUEBA ==============================================
TEMPLATE = r"C:\Users\Hector\Downloads\Certificaddo FOCA VIOLENCIA SEXUAL.docx"

# Datos reales de la persona (de PLANTA FOCA.xlsx)
NOMBRE     = "ACOSTA MORELO ANDREA DEL CARMEN"
DOC_NUM    = "1045737800"
CIUDAD     = "BARRANQUILLA"
FECHA      = "15 de marzo de 2026"
FECHA_CORTA = "15/03/2026"
VALIDO_HASTA = "15/03/2028"
HORAS      = "4 horas"
REGISTRO_ID = f"FOCA-{DOC_NUM}"

# URL del portal — el QR precargará esta cédula automáticamente
PORTAL_BASE = "https://deineracosta.github.io/certificados-sst/"
QR_URL      = f"{PORTAL_BASE}?cc={DOC_NUM}"

OUTPUT_DIR = Path(__file__).parent / "pruebas"
OUTPUT_DIR.mkdir(exist_ok=True)


# ==== HELPERS =================================================================

def generar_qr(url: str, out_path: Path) -> None:
    qr = qrcode.QRCode(
        version=None,
        error_correction=ERROR_CORRECT_H,
        box_size=15,
        border=2,
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="#00467F", back_color="white")
    img.save(out_path)


def reemplazar_en_paragrafo(para, mapping: dict) -> bool:
    changed = False
    texto = para.text
    nuevo = texto
    for viejo, nuevo_val in mapping.items():
        if viejo in nuevo:
            nuevo = nuevo.replace(viejo, nuevo_val)
            changed = True
    if changed and para.runs:
        para.runs[0].text = nuevo
        for r in para.runs[1:]:
            r.text = ""
    return changed


def in_to_emu(inches: float) -> int:
    return int(inches * 914400)


def add_floating_image(doc, image_path: str, width_in: float,
                       pos_from_left_in: float, pos_from_top_in: float,
                       image_id: int = 100):
    """
    Agrega una imagen FLOTANTE anclada a posición absoluta en la página.
    No afecta el flujo del texto (wrap = none).
    """
    # Paso 1: agregar como inline en un párrafo temporal para obtener el rId
    # (usaremos el último párrafo existente del documento)
    last_para = doc.paragraphs[-1]
    run = last_para.add_run()
    run.add_picture(image_path, width=Inches(width_in))

    # Paso 2: extraer el drawing recién insertado
    drawing = run._element.find(qn('w:drawing'))
    inline = drawing.find(qn('wp:inline'))

    # Obtener el rId del blip
    blip = inline.find('.//' + qn('a:blip'))
    r_embed = blip.get(qn('r:embed'))

    # Extensiones (tamaño en EMU)
    extent = inline.find(qn('wp:extent'))
    cx = extent.get('cx')
    cy = extent.get('cy')

    # Paso 3: construir el <wp:anchor> con posición absoluta
    x_emu = in_to_emu(pos_from_left_in)
    y_emu = in_to_emu(pos_from_top_in)

    anchor_xml = f'''<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
       xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
       xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
       xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:drawing>
    <wp:anchor distT="0" distB="0" distL="0" distR="0"
               simplePos="0" relativeHeight="251659264" behindDoc="0"
               locked="0" layoutInCell="1" allowOverlap="1">
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="page">
        <wp:posOffset>{x_emu}</wp:posOffset>
      </wp:positionH>
      <wp:positionV relativeFrom="page">
        <wp:posOffset>{y_emu}</wp:posOffset>
      </wp:positionV>
      <wp:extent cx="{cx}" cy="{cy}"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
      <wp:docPr id="{image_id}" name="QR-Verificacion"/>
      <wp:cNvGraphicFramePr>
        <a:graphicFrameLocks noChangeAspect="1"/>
      </wp:cNvGraphicFramePr>
      <a:graphic>
        <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
          <pic:pic>
            <pic:nvPicPr>
              <pic:cNvPr id="0" name="qr.png"/>
              <pic:cNvPicPr/>
            </pic:nvPicPr>
            <pic:blipFill>
              <a:blip r:embed="{r_embed}"/>
              <a:stretch><a:fillRect/></a:stretch>
            </pic:blipFill>
            <pic:spPr>
              <a:xfrm>
                <a:off x="0" y="0"/>
                <a:ext cx="{cx}" cy="{cy}"/>
              </a:xfrm>
              <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
            </pic:spPr>
          </pic:pic>
        </a:graphicData>
      </a:graphic>
    </wp:anchor>
  </w:drawing>
</w:r>'''

    # Reemplazar el run inline por el anchor
    new_run = parse_xml(anchor_xml)
    run._element.getparent().replace(run._element, new_run)


def add_floating_text(doc, text: str,
                      pos_from_left_in: float, pos_from_top_in: float,
                      width_in: float, height_in: float,
                      font_size: int = 9, italic: bool = True,
                      color_hex: str = "555555",
                      text_id: int = 101,
                      bold: bool = False,
                      align: str = "center",
                      letter_spacing: int = 0):
    """Agrega un cuadro de texto flotante anclado a posición absoluta.

    align: 'center' | 'left' | 'right'
    letter_spacing: en 1/20 de punto (10 = 0.5pt)
    """
    last_para = doc.paragraphs[-1]

    x_emu = in_to_emu(pos_from_left_in)
    y_emu = in_to_emu(pos_from_top_in)
    w_emu = in_to_emu(width_in)
    h_emu = in_to_emu(height_in)

    tx_color = color_hex.upper()
    spacing_xml = f'<w:spacing w:val="{letter_spacing}"/>' if letter_spacing else ''

    anchor_xml = f'''<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
       xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
       xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
  <w:drawing>
    <wp:anchor distT="0" distB="0" distL="0" distR="0"
               simplePos="0" relativeHeight="{251659265 + text_id}" behindDoc="0"
               locked="0" layoutInCell="1" allowOverlap="1">
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="page">
        <wp:posOffset>{x_emu}</wp:posOffset>
      </wp:positionH>
      <wp:positionV relativeFrom="page">
        <wp:posOffset>{y_emu}</wp:posOffset>
      </wp:positionV>
      <wp:extent cx="{w_emu}" cy="{h_emu}"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
      <wp:docPr id="{text_id}" name="TextBox-{text_id}"/>
      <wp:cNvGraphicFramePr/>
      <a:graphic>
        <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
          <wps:wsp>
            <wps:cNvSpPr txBox="1"/>
            <wps:spPr>
              <a:xfrm>
                <a:off x="0" y="0"/>
                <a:ext cx="{w_emu}" cy="{h_emu}"/>
              </a:xfrm>
              <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
              <a:noFill/>
              <a:ln><a:noFill/></a:ln>
            </wps:spPr>
            <wps:txbx>
              <w:txbxContent>
                <w:p>
                  <w:pPr>
                    <w:jc w:val="{align}"/>
                    <w:spacing w:before="0" w:after="0"/>
                  </w:pPr>
                  <w:r>
                    <w:rPr>
                      <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                      <w:sz w:val="{font_size * 2}"/>
                      {"<w:i/>" if italic else ""}
                      {"<w:b/>" if bold else ""}
                      {spacing_xml}
                      <w:color w:val="{tx_color}"/>
                    </w:rPr>
                    <w:t xml:space="preserve">{text}</w:t>
                  </w:r>
                </w:p>
              </w:txbxContent>
            </wps:txbx>
            <wps:bodyPr wrap="square" anchor="ctr" lIns="0" tIns="0" rIns="0" bIns="0"/>
          </wps:wsp>
        </a:graphicData>
      </a:graphic>
    </wp:anchor>
  </w:drawing>
</w:r>'''

    new_run = parse_xml(anchor_xml)
    last_para._element.append(new_run)


def add_floating_line(doc, pos_x_in: float, pos_y_in: float,
                      width_in: float, thickness_pt: float = 0.75,
                      color_hex: str = "CCCCCC", line_id: int = 500):
    """Agrega una línea horizontal fina, útil como separador."""
    last_para = doc.paragraphs[-1]
    x_emu = in_to_emu(pos_x_in)
    y_emu = in_to_emu(pos_y_in)
    w_emu = in_to_emu(width_in)
    h_emu = int(thickness_pt * 12700)  # pt -> EMU

    anchor_xml = f'''<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
       xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
       xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
  <w:drawing>
    <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
               relativeHeight="{251659265 + line_id}" behindDoc="0"
               locked="0" layoutInCell="1" allowOverlap="1">
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="page"><wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>
      <wp:positionV relativeFrom="page"><wp:posOffset>{y_emu}</wp:posOffset></wp:positionV>
      <wp:extent cx="{w_emu}" cy="{h_emu}"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
      <wp:docPr id="{line_id}" name="Sep-{line_id}"/>
      <wp:cNvGraphicFramePr/>
      <a:graphic>
        <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
          <wps:wsp>
            <wps:cNvSpPr/>
            <wps:spPr>
              <a:xfrm><a:off x="0" y="0"/><a:ext cx="{w_emu}" cy="{h_emu}"/></a:xfrm>
              <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
              <a:solidFill><a:srgbClr val="{color_hex.upper()}"/></a:solidFill>
              <a:ln><a:noFill/></a:ln>
            </wps:spPr>
            <wps:bodyPr/>
          </wps:wsp>
        </a:graphicData>
      </a:graphic>
    </wp:anchor>
  </w:drawing>
</w:r>'''
    last_para._element.append(parse_xml(anchor_xml))


# ==== MAIN ====================================================================

def main() -> None:
    print("\n[FOCA] Generando certificado de PRUEBA v2 (QR flotante)...\n")

    # 1. Generar el QR
    qr_path = OUTPUT_DIR / f"qr-{DOC_NUM}.png"
    generar_qr(QR_URL, qr_path)
    print(f"  [OK] QR generado: {qr_path.name}")

    # 2. Abrir plantilla y reemplazar texto
    doc = Document(TEMPLATE)

    reemplazos = {
        "LUIS JOSE ESCAF JARABA": NOMBRE,
        "6813307": DOC_NUM,
        "20 de enero de 2026": FECHA,
        "Curso de Violencia Sexual, realizado el":
            f"Curso de Violencia Sexual, realizado en {CIUDAD} el",
        "una intensidad horaria de 4 horas": f"una intensidad horaria de {HORAS}",
    }

    cambios = 0
    for para in doc.paragraphs:
        if reemplazar_en_paragrafo(para, reemplazos):
            cambios += 1
    print(f"  [OK] {cambios} parrafos modificados")

    # 3. QR flotante en esquina inferior derecha
    add_floating_image(
        doc,
        str(qr_path),
        width_in=1.25,
        pos_from_left_in=9.35,
        pos_from_top_in=6.85,
        image_id=100,
    )
    print("  [OK] QR flotante posicionado en esquina inf-derecha")

    # 4. Barra de metadatos estilo CAEM (INTENSIDAD | FECHA | VÁLIDO HASTA | ID)
    # Página: 11" ancho × 8.5" alto (landscape)
    # Fila de metadatos abajo, antes del QR (que está en x=9.35)

    # Línea separadora fina arriba de la barra
    add_floating_line(doc, pos_x_in=0.7, pos_y_in=7.15,
                      width_in=8.5, thickness_pt=0.5,
                      color_hex="D0D0D0", line_id=500)

    # 4 columnas de metadata
    metadata = [
        ("INTENSIDAD HORARIA", HORAS.upper()),
        ("FECHA DE REALIZACIÓN", FECHA_CORTA),
        ("VÁLIDO HASTA", VALIDO_HASTA),
        ("ID DE VERIFICACIÓN", REGISTRO_ID),
    ]

    col_width = 2.1
    start_x = 0.7
    y_label = 7.30
    y_value = 7.55

    for i, (label, value) in enumerate(metadata):
        x = start_x + i * col_width
        # Label (pequeño, gris, uppercase con letter-spacing)
        add_floating_text(
            doc, label,
            pos_from_left_in=x,
            pos_from_top_in=y_label,
            width_in=col_width - 0.1,
            height_in=0.22,
            font_size=7,
            italic=False,
            bold=False,
            color_hex="808080",
            letter_spacing=20,  # 1pt letter-spacing
            text_id=200 + i * 2,
            align="center",
        )
        # Value (más grande, azul FOCA, bold)
        add_floating_text(
            doc, value,
            pos_from_left_in=x,
            pos_from_top_in=y_value,
            width_in=col_width - 0.1,
            height_in=0.32,
            font_size=12,
            italic=False,
            bold=True,
            color_hex="00467F",
            text_id=201 + i * 2,
            align="center",
        )

    print("  [OK] Barra de metadatos estilo CAEM agregada")

    # 5. Guardar
    nombre_archivo = f"PRUEBA-v2-{DOC_NUM}-{NOMBRE.replace(' ', '_')}.docx"
    output_path = OUTPUT_DIR / nombre_archivo
    doc.save(output_path)

    print()
    print("=== LISTO ===")
    print(f"  Archivo: {output_path}")
    print(f"  QR apunta a: {QR_URL}")
    print()


if __name__ == "__main__":
    main()
