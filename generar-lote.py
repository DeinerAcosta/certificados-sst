"""
FOCA · Generador de LOTE de certificados
==========================================

Lee las N primeras personas del Excel PLANTA FOCA y para cada una:
  · Genera su PDF (usando la plantilla FOCA + QR flotante estilo CAEM)
  · Agrega su entrada al datos.js

Configurá abajo:
  · CANTIDAD  → cuántas personas del Excel procesar
  · CIUDAD    → ciudad de la capacitación
  · FECHA     → fecha real de la capacitación
  · CURSO     → título del curso
  · HORAS     → intensidad horaria
"""

from pathlib import Path
import json
import qrcode
from qrcode.constants import ERROR_CORRECT_H
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import openpyxl

# ==== CONFIGURACIÓN ===========================================================
CANTIDAD    = 10                    # ← primeras N personas del Excel
CIUDAD      = "BARRANQUILLA"
FECHA       = "15 de marzo de 2026"
FECHA_CORTA = "15/03/2026"
VALIDO_HASTA = "15 de marzo de 2028"  # 2 años por defecto
VALIDO_HASTA_CORTA = "15/03/2028"
CURSO       = "Violencia Sexual"
HORAS       = "4 horas"

EXCEL       = r"C:\Users\Hector\Downloads\PLANTA FOCA.xlsx"
TEMPLATE    = r"C:\Users\Hector\Downloads\Certificaddo FOCA VIOLENCIA SEXUAL.docx"

# URL pública donde vive el portal
PORTAL_BASE = "https://deineracosta.github.io/certificados-sst/"

# ==== SALIDA ==================================================================
BASE_DIR    = Path(__file__).parent
PDFS_DIR    = BASE_DIR / "pdfs"
DATOS_JS    = BASE_DIR / "datos.js"
PDFS_DIR.mkdir(exist_ok=True)


# ==== HELPERS =================================================================

def in_to_emu(inches: float) -> int:
    return int(inches * 914400)


def generar_qr(url: str, out_path: Path) -> None:
    qr = qrcode.QRCode(
        version=None,
        error_correction=ERROR_CORRECT_H,
        box_size=15,
        border=2,
    )
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color="#00467F", back_color="white")
    img.save(out_path)


def reemplazar_en_paragrafo(para, mapping: dict) -> bool:
    changed = False
    texto = para.text
    nuevo = texto
    for viejo, val in mapping.items():
        if viejo in nuevo:
            nuevo = nuevo.replace(viejo, val)
            changed = True
    if changed and para.runs:
        para.runs[0].text = nuevo
        for r in para.runs[1:]:
            r.text = ""
    return changed


def add_floating_image(doc, image_path, width_in, pos_x_in, pos_y_in, image_id):
    last_para = doc.paragraphs[-1]
    run = last_para.add_run()
    run.add_picture(image_path, width=Inches(width_in))
    drawing = run._element.find(qn('w:drawing'))
    inline = drawing.find(qn('wp:inline'))
    blip = inline.find('.//' + qn('a:blip'))
    r_embed = blip.get(qn('r:embed'))
    extent = inline.find(qn('wp:extent'))
    cx, cy = extent.get('cx'), extent.get('cy')

    x = in_to_emu(pos_x_in)
    y = in_to_emu(pos_y_in)

    anchor_xml = f'''<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
       xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
       xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
       xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <w:drawing>
    <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
               relativeHeight="{251659264 + image_id}" behindDoc="0"
               locked="0" layoutInCell="1" allowOverlap="1">
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="page"><wp:posOffset>{x}</wp:posOffset></wp:positionH>
      <wp:positionV relativeFrom="page"><wp:posOffset>{y}</wp:posOffset></wp:positionV>
      <wp:extent cx="{cx}" cy="{cy}"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
      <wp:docPr id="{image_id}" name="QR-{image_id}"/>
      <wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect="1"/></wp:cNvGraphicFramePr>
      <a:graphic>
        <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
          <pic:pic>
            <pic:nvPicPr><pic:cNvPr id="0" name="qr.png"/><pic:cNvPicPr/></pic:nvPicPr>
            <pic:blipFill><a:blip r:embed="{r_embed}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>
            <pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>
              <a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr>
          </pic:pic>
        </a:graphicData>
      </a:graphic>
    </wp:anchor>
  </w:drawing>
</w:r>'''
    new_run = parse_xml(anchor_xml)
    run._element.getparent().replace(run._element, new_run)


def add_floating_text(doc, text, pos_x_in, pos_y_in, width_in, height_in,
                      font_size, italic, color_hex, text_id, bold=False,
                      align="center", letter_spacing=0):
    last_para = doc.paragraphs[-1]
    x, y = in_to_emu(pos_x_in), in_to_emu(pos_y_in)
    w, h = in_to_emu(width_in), in_to_emu(height_in)
    spacing_xml = f'<w:spacing w:val="{letter_spacing}"/>' if letter_spacing else ''
    anchor_xml = f'''<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
       xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
       xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
  <w:drawing>
    <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
               relativeHeight="{251659300 + text_id}" behindDoc="0"
               locked="0" layoutInCell="1" allowOverlap="1">
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="page"><wp:posOffset>{x}</wp:posOffset></wp:positionH>
      <wp:positionV relativeFrom="page"><wp:posOffset>{y}</wp:posOffset></wp:positionV>
      <wp:extent cx="{w}" cy="{h}"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
      <wp:docPr id="{text_id}" name="Txt-{text_id}"/>
      <wp:cNvGraphicFramePr/>
      <a:graphic>
        <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
          <wps:wsp>
            <wps:cNvSpPr txBox="1"/>
            <wps:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="{w}" cy="{h}"/></a:xfrm>
              <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
              <a:noFill/><a:ln><a:noFill/></a:ln></wps:spPr>
            <wps:txbx><w:txbxContent>
              <w:p>
                <w:pPr><w:jc w:val="{align}"/><w:spacing w:before="0" w:after="0"/></w:pPr>
                <w:r>
                  <w:rPr>
                    <w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>
                    <w:sz w:val="{font_size * 2}"/>
                    {"<w:i/>" if italic else ""}
                    {"<w:b/>" if bold else ""}
                    {spacing_xml}
                    <w:color w:val="{color_hex.upper()}"/>
                  </w:rPr>
                  <w:t xml:space="preserve">{text}</w:t>
                </w:r>
              </w:p>
            </w:txbxContent></wps:txbx>
            <wps:bodyPr wrap="square" anchor="ctr" lIns="0" tIns="0" rIns="0" bIns="0"/>
          </wps:wsp>
        </a:graphicData>
      </a:graphic>
    </wp:anchor>
  </w:drawing>
</w:r>'''
    last_para._element.append(parse_xml(anchor_xml))


def add_floating_line(doc, pos_x_in, pos_y_in, width_in,
                      thickness_pt=0.5, color_hex="D0D0D0", line_id=500):
    last_para = doc.paragraphs[-1]
    x, y = in_to_emu(pos_x_in), in_to_emu(pos_y_in)
    w = in_to_emu(width_in)
    h = int(thickness_pt * 12700)
    anchor_xml = f'''<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
       xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
       xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
  <w:drawing>
    <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
               relativeHeight="{251659265 + line_id}" behindDoc="0"
               locked="0" layoutInCell="1" allowOverlap="1">
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="page"><wp:posOffset>{x}</wp:posOffset></wp:positionH>
      <wp:positionV relativeFrom="page"><wp:posOffset>{y}</wp:posOffset></wp:positionV>
      <wp:extent cx="{w}" cy="{h}"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
      <wp:docPr id="{line_id}" name="Sep-{line_id}"/>
      <wp:cNvGraphicFramePr/>
      <a:graphic>
        <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
          <wps:wsp>
            <wps:cNvSpPr/>
            <wps:spPr>
              <a:xfrm><a:off x="0" y="0"/><a:ext cx="{w}" cy="{h}"/></a:xfrm>
              <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
              <a:solidFill><a:srgbClr val="{color_hex.upper()}"/></a:solidFill>
              <a:ln><a:noFill/></a:ln>
            </wps:spPr>
            <wps:bodyPr/>
          </wps:wsp>
        </a:graphicData>
      </a:graphic>
    </wp:anchor>
  </w:drawing>
</w:r>'''
    last_para._element.append(parse_xml(anchor_xml))


def generar_certificado(nombre: str, cedula: str, tipo_doc: str, cargo: str) -> Path:
    """Genera el PDF (bueno, .docx) para una persona. Devuelve la ruta."""
    qr_url = f"{PORTAL_BASE}?cc={cedula}"
    qr_path = PDFS_DIR / f"_qr-{cedula}.png"
    generar_qr(qr_url, qr_path)

    doc = Document(TEMPLATE)
    reemplazos = {
        "LUIS JOSE ESCAF JARABA": nombre,
        "6813307": cedula,
        "20 de enero de 2026": FECHA,
        f"Curso de {CURSO}, realizado el" if CURSO != "Violencia Sexual"
            else "Curso de Violencia Sexual, realizado el":
            f"Curso de {CURSO}, realizado en {CIUDAD} el",
        "una intensidad horaria de 4 horas": f"una intensidad horaria de {HORAS}",
    }
    for para in doc.paragraphs:
        reemplazar_en_paragrafo(para, reemplazos)

    # QR flotante abajo derecha
    add_floating_image(doc, str(qr_path), 1.25, 9.35, 6.85, image_id=100)

    # Barra de metadatos estilo CAEM
    add_floating_line(doc, pos_x_in=0.7, pos_y_in=7.15,
                      width_in=8.5, thickness_pt=0.5,
                      color_hex="D0D0D0", line_id=500)

    metadata = [
        ("INTENSIDAD HORARIA", HORAS.upper()),
        ("FECHA DE REALIZACIÓN", FECHA_CORTA),
        ("VÁLIDO HASTA", VALIDO_HASTA_CORTA),
        ("ID DE VERIFICACIÓN", f"FOCA-{cedula}"),
    ]
    col_width = 2.1
    start_x = 0.7
    for i, (label, value) in enumerate(metadata):
        x = start_x + i * col_width
        add_floating_text(doc, label, x, 7.30, col_width - 0.1, 0.22,
                          font_size=7, italic=False, color_hex="808080",
                          text_id=200 + i * 2, letter_spacing=20)
        add_floating_text(doc, value, x, 7.55, col_width - 0.1, 0.32,
                          font_size=12, italic=False, color_hex="00467F",
                          text_id=201 + i * 2, bold=True)

    docx_path = PDFS_DIR / f"{cedula}.docx"
    doc.save(docx_path)
    qr_path.unlink()  # limpiar QR temporal
    return docx_path


def leer_personas_del_excel(n: int) -> list[dict]:
    """Lee las primeras n personas del Excel PLANTA FOCA."""
    wb = openpyxl.load_workbook(EXCEL, data_only=True)
    ws = wb.active
    personas = []
    for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        if len(personas) >= n:
            break
        tipo_doc, num_doc, nombre, cargo = row[:4]
        if not (tipo_doc and num_doc and nombre):
            continue
        personas.append({
            "tipo_doc": str(tipo_doc).strip(),
            "cedula": str(num_doc).strip(),
            "nombre": str(nombre).strip(),
            "cargo": str(cargo).strip() if cargo else "-",
        })
    return personas


def escribir_datos_js(personas: list[dict]) -> None:
    """Escribe el datos.js que el portal lee."""
    db = {}
    for p in personas:
        db[p["cedula"]] = {
            "nombre": p["nombre"],
            "documento": f"{p['tipo_doc']} {p['cedula']}",
            "cargo": p["cargo"].title(),
            "empresa": "FOCA",
            "certificados": [
                {
                    "titulo": f"Curso de {CURSO}",
                    "fecha": FECHA,
                    "valido_hasta": VALIDO_HASTA,
                    "horas": HORAS,
                    "pdf": f"pdfs/{p['cedula']}.pdf",
                }
            ]
        }
    contenido = (
        "// Base de datos del portal — editá este archivo para agregar personas\n"
        "// Después de editar: git add + git commit + git push\n\n"
        f"window.PERSONAS = {json.dumps(db, ensure_ascii=False, indent=2)};\n"
    )
    DATOS_JS.write_text(contenido, encoding="utf-8")


# ==== MAIN ====================================================================

def main() -> None:
    print(f"\n[FOCA] Lote de {CANTIDAD} certificados de prueba\n")
    print(f"  Ciudad:  {CIUDAD}")
    print(f"  Fecha:   {FECHA}")
    print(f"  Curso:   {CURSO} ({HORAS})")
    print(f"  Portal:  {PORTAL_BASE}\n")

    personas = leer_personas_del_excel(CANTIDAD)
    print(f"Encontradas {len(personas)} personas en Excel.\n")

    for i, p in enumerate(personas, 1):
        path = generar_certificado(p["nombre"], p["cedula"], p["tipo_doc"], p["cargo"])
        print(f"  [{i:2d}/{len(personas)}] {p['cedula']} - {p['nombre'][:40]}")

    escribir_datos_js(personas)
    print(f"\n[OK] datos.js escrito: {DATOS_JS}")
    print(f"[OK] {len(personas)} DOCX generados en: {PDFS_DIR}")
    print("\nSIGUIENTE: convertir DOCX -> PDF con Word (PowerShell)")


if __name__ == "__main__":
    main()
